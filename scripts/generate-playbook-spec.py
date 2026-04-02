#!/usr/bin/env python3
"""Generate the Agent Playbook Functional Specification as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1a, 0x15, 0x35)

# Helper functions
def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_checklist(items, indent=0):
    """Add checklist items as bullet points."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.27 + indent * 1.27)
        p.paragraph_format.space_after = Pt(2)
        run = p.runs[0] if p.runs else p.add_run(item)
        if not p.runs[0].text:
            p.runs[0].text = item
        run.font.size = Pt(10)

def section(title, level=2):
    doc.add_heading(title, level=level)

def para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bold_para(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p

# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Agent Playbook', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Functional Specification', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('ListingFlow Real Estate CRM\n').bold = True
meta.add_run('Task Execution Framework for AI Development Agents\n\n')
meta.add_run('Version: 1.0\n')
meta.add_run('Date: March 2026\n')
meta.add_run('Status: Draft → Review\n\n')
meta.add_run('Classification: Internal Developer Tooling\n')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Problem Statement & Goals',
    '3. Personas & Actors',
    '4. System Architecture',
    '5. Task Classification Engine',
    '6. Mandatory Pre-Flight Protocol',
    '7. Task Type Playbooks (10 Types)',
    '   7.1  CODING (feature / bugfix / refactor / script)',
    '   7.2  TESTING (unit / integration / e2e / eval)',
    '   7.3  DEBUGGING (error / performance / data_issue)',
    '   7.4  DESIGN_SPEC (architecture / feature / api / migration)',
    '   7.5  RAG_KB (pipeline / tuning / evaluation / content)',
    '   7.6  ORCHESTRATION (workflow / trigger / pipeline / agent)',
    '   7.7  INTEGRATION (api_connect / webhook / auth / data_sync)',
    '   7.8  DOCS (spec / guide / runbook / changelog)',
    '   7.9  EVAL (metrics / golden_set / ab_test / quality_gate)',
    '   7.10 INFO_QA (explain / compare / recommend)',
    '8. Post-Task Validation Protocol',
    '9. Model Chaining Strategy',
    '10. Decision Framework: Ask vs Act',
    '11. Existing Infrastructure Mapping',
    '12. Integration Points & Non-Disruption Guarantees',
    '13. Test Scenarios for the Playbook Itself',
    '14. Appendix: Quick Reference Card',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

para('This document specifies a Task Execution Framework (the "Agent Playbook") that governs how AI development agents (Claude Code) operate on the ListingFlow CRM codebase. The playbook is NOT application code — it is a set of operational rules, checklists, and decision frameworks that the AI agent must follow when developers assign it tasks.')

para('The playbook addresses a critical gap: today, the AI agent jumps directly into execution without systematically classifying the task, loading context, or following task-specific quality checklists. This leads to inconsistent quality, missed edge cases, and unnecessary rework.')

bold_para('Key Outcomes:')
add_checklist([
    'Every task is classified before execution begins (10 task types × 2-4 subtypes)',
    'Each task type has a mandatory, ordered checklist that the agent follows',
    'Existing infrastructure is preserved — the playbook layers ON TOP, never replaces',
    'Model chaining is formalized (Haiku for routing, Sonnet for coding, Opus for architecture)',
    'Post-task validation is mandatory (tests, lint, self-review)',
    'All checklists cover every possible scenario (happy path, edge cases, error conditions, concurrent operations)',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT & GOALS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Problem Statement & Goals', level=1)

section('2.1 Problems Observed')
add_table(
    ['#', 'Problem', 'Impact', 'Frequency'],
    [
        ['P1', 'Agent starts coding without analyzing scope or affected files', 'Unintended side effects, missed integration points, broken features', 'Every coding task'],
        ['P2', 'No systematic debugging methodology — agent tries random fixes', 'Slower resolution, band-aid fixes that mask root cause', 'Every bug report'],
        ['P3', 'Design specs produced in inconsistent formats', 'Hard to review, missing sections, incomplete analysis', 'Every design task'],
        ['P4', 'Agent does not self-check code after writing it', 'Logic errors, unused variables, unhandled branches slip through', 'Frequent'],
        ['P5', 'No task classification — agent guesses the approach', 'Wrong methodology applied (e.g., refactoring when debugging needed)', 'Occasional'],
        ['P6', 'Testing checklists incomplete — happy path only', 'Edge cases, error conditions, concurrent scenarios missed', 'Every testing task'],
        ['P7', 'Model chaining is ad-hoc — no per-phase guidance', 'Expensive models used for simple tasks, cheap models for complex ones', 'Ongoing'],
        ['P8', 'Multiple developers working — agent doesn\'t check for conflicts', 'Disrupts other developers\' in-progress work', 'Risk per session'],
    ]
)

section('2.2 Goals')
add_table(
    ['#', 'Goal', 'Success Metric'],
    [
        ['G1', 'Every task classified before execution', '100% of tasks have explicit classification in agent output'],
        ['G2', 'Per-task checklists followed completely', 'Zero skipped phases (auditable via agent output)'],
        ['G3', 'All test scenarios cover edge cases and error conditions', 'Every TESTING task produces ≥5 edge case tests'],
        ['G4', 'Existing code/infra never disrupted', 'Zero regressions introduced by playbook adoption'],
        ['G5', 'Model chaining cost-optimized', 'Haiku used for ≥60% of sub-tasks, Opus only for architecture'],
        ['G6', 'Self-check catches errors before user review', '≥80% of issues caught by agent before presenting output'],
    ]
)

section('2.3 Non-Goals')
add_checklist([
    'This playbook does NOT change any application code, database schema, or API contracts',
    'This playbook does NOT replace existing CLAUDE.md coding conventions — it references them',
    'This playbook does NOT apply to end-user interactions (CRM chatbot has its own guardrails)',
    'This playbook does NOT require any npm packages, migrations, or deployments',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. PERSONAS & ACTORS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Personas & Actors', level=1)

add_table(
    ['Actor', 'Role', 'Interaction with Playbook'],
    [
        ['Developer (Primary)', 'Assigns tasks to Claude Code via natural language', 'Provides task requests; reviews agent output; gives feedback'],
        ['Claude Code Agent (Executor)', 'AI assistant executing developer tasks', 'Classifies task → loads playbook → executes checklist → validates'],
        ['Sub-Agents (Haiku/Sonnet)', 'Delegated workers for specific phases', 'Perform file search, code writing, test execution per model chain rules'],
        ['CI/CD System', 'Automated build/test/deploy pipeline', 'Validates agent output via test suite, lint, build checks'],
        ['Other Developers', 'Team members working on same codebase', 'Their work must not be disrupted — agent checks for conflicts'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. System Architecture', level=1)

section('4.1 Execution Flow')
para('''The agent follows a strict 4-step loop for every task:

Step 0: Pre-Flight → Verify environment, check git status, load memory
Step 1: Classify → Determine task_type:subtype with confidence + reasoning
Step 2: Execute → Follow the per-task checklist phase by phase
Step 3: Validate → Run tests, self-check, verify no regressions

Each step MUST complete before the next begins. The agent cannot skip steps.''')

section('4.2 Component Diagram')
para('''
┌─────────────────────────────────────────────────────────┐
│                    DEVELOPER REQUEST                      │
└─────────────────────┬─────────────────────────────────────┘
                      │
        ┌─────────────▼──────────────┐
        │   STEP 0: PRE-FLIGHT       │
        │   • Health check            │
        │   • Git status (branch,     │
        │     conflicts, stash)       │
        │   • Load memory rules       │
        └─────────────┬──────────────┘
                      │
        ┌─────────────▼──────────────┐
        │   STEP 1: CLASSIFY          │
        │   • Task type + subtype     │
        │   • Affected components     │
        │   • Confidence + reasoning  │
        │   • Ask if confidence low   │
        └─────────────┬──────────────┘
                      │
        ┌─────────────▼──────────────┐
        │   STEP 2: EXECUTE           │
        │   • Load playbook checklist │
        │   • Execute phase by phase  │
        │   • Record artifacts        │
        │   • Model chain per phase   │
        └─────────────┬──────────────┘
                      │
        ┌─────────────▼──────────────┐
        │   STEP 3: VALIDATE          │
        │   • Run tests               │
        │   • Self-check code         │
        │   • Verify no regressions   │
        │   • Update todo list        │
        └────────────────────────────┘
''')

section('4.3 Where the Playbook Lives')
add_table(
    ['Component', 'Location', 'Format'],
    [
        ['Playbook definition', '.claude/agent-playbook.md', 'Markdown (loaded as project instructions)'],
        ['Existing coding conventions', 'CLAUDE.md', 'Referenced, not duplicated'],
        ['Behavioral rules', 'Memory files (feedback_*.md)', 'Referenced, not duplicated'],
        ['Deploy skill', '.claude/skills/deploy.md', 'Unchanged — playbook references it'],
        ['Test skill', '.claude/skills/test.md', 'Unchanged — playbook references it'],
        ['Health check', 'scripts/health-check.sh', 'Unchanged — playbook calls it'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. TASK CLASSIFICATION ENGINE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Task Classification Engine', level=1)

section('5.1 Classification Output Format')
para('''Before any execution, the agent MUST output:

Task Type: CODING:feature
Confidence: high | medium | low
Reasoning: [1-2 sentence explanation of why this classification]
Affected: [list of files, tables, APIs, components]
Secondary types: [if task spans multiple types, list them]''')

section('5.2 Task Type Taxonomy')
add_table(
    ['Type', 'Subtypes', 'Trigger Phrases (Examples)', 'Model Chain Default'],
    [
        ['CODING', 'feature, bugfix, refactor, script', '"Add a new field", "Implement X", "Build the Y page", "Write a migration"', 'Haiku→classify, Sonnet→code, Opus→review complex'],
        ['TESTING', 'unit, integration, e2e, eval', '"Write tests for", "Run the test suite", "Check coverage", "Validate"', 'Sonnet→plan+write, Haiku→run'],
        ['DEBUGGING', 'error, performance, data_issue', '"Fix this error", "Why is X broken", "This returns wrong data", stack traces', 'Opus→hypothesize, Sonnet→inspect+fix'],
        ['DESIGN_SPEC', 'architecture, feature_spec, api_design, migration_plan', '"Design a system for", "Write a spec", "Plan the architecture", "How should we build"', 'Opus→design, Sonnet→details'],
        ['RAG_KB', 'pipeline, tuning, evaluation, content', '"Improve retrieval", "Add a new data source", "Tune the embeddings", "Evaluate RAG quality"', 'Opus→strategy, Sonnet→implement'],
        ['ORCHESTRATION', 'workflow, trigger, pipeline, agent_config', '"Add a workflow trigger", "Modify the AI pipeline", "Change trust level logic"', 'Opus→design, Sonnet→implement'],
        ['INTEGRATION', 'api_connect, webhook, auth, data_sync', '"Connect to X API", "Add webhook handler", "Integrate with Y service"', 'Sonnet→implement, Haiku→test'],
        ['DOCS', 'spec, guide, runbook, changelog', '"Document X", "Write a guide for", "Update the README", "Create a runbook"', 'Sonnet→write, Haiku→verify links'],
        ['EVAL', 'metrics, golden_set, ab_test, quality_gate', '"Define success metrics", "Create test scenarios", "Evaluate performance", "Score quality"', 'Opus→define, Sonnet→execute'],
        ['INFO_QA', 'explain, compare, recommend', '"Explain how X works", "Compare A vs B", "What should we use for", "How does the codebase handle"', 'Haiku→search, Sonnet→synthesize'],
    ]
)

section('5.3 Classification Rules')
add_checklist([
    'If the request involves MULTIPLE types, pick the PRIMARY and note secondaries',
    'If the request is a compound task ("build X and write tests"), classify the primary (CODING) and note TESTING as secondary',
    'If confidence is LOW (request is ambiguous), ask ONE clarifying question — do not guess',
    'If the request says "make it better" without specifics, ask what "better" means',
    'Always identify AFFECTED COMPONENTS during classification (files, tables, APIs)',
    'If the request mentions an error or stack trace → always classify as DEBUGGING first (not CODING:bugfix)',
    'If the request mentions "plan" or "design" before "build" → classify as DESIGN_SPEC first',
    'If the request says "refactor" → classify as CODING:refactor (NOT CODING:feature)',
])

section('5.4 Compound Task Handling')
para('When a single request spans multiple task types:')
add_checklist([
    'Classify the primary type and execute its full checklist first',
    'Then execute secondary type checklists in order',
    'Never interleave — complete one before starting the next',
    'Example: "Add a contact field and write tests" → CODING:feature (full checklist) → TESTING:unit (full checklist)',
    'Example: "Debug this error and fix it" → DEBUGGING:error (full checklist including fix proposal) → CODING:bugfix (if fix needs implementation)',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. MANDATORY PRE-FLIGHT PROTOCOL
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Mandatory Pre-Flight Protocol', level=1)

para('Before ANY task execution — no exceptions:')

section('6.1 Environment Verification')
add_checklist([
    'Run: bash scripts/health-check.sh (from memory: feedback_auto_health_check.md)',
    'Verify output shows no critical failures (build, env vars, DB connectivity)',
    'If health check fails: fix the environment issue BEFORE doing any task work',
])

section('6.2 Git Status Check')
add_checklist([
    'Confirm current branch is dev (never work on main)',
    'Check for uncommitted changes (may be another developer\'s work — DO NOT discard)',
    'Check for unmerged changes from remote (pull if behind)',
    'If there are conflicts or stashed changes: inform the developer before proceeding',
])

section('6.3 Memory & Context Loading')
add_checklist([
    'Read MEMORY.md for behavioral rules and project context',
    'Check for any recent memory entries relevant to the current task',
    'Load any task-specific notes from previous sessions',
])

section('6.4 Pre-Flight Failure Scenarios')
add_table(
    ['Scenario', 'Action', 'Do NOT'],
    [
        ['Health check shows build failure', 'Report to developer, offer to fix build first', 'Start the requested task with a broken build'],
        ['Branch is main', 'Switch to dev: git checkout dev', 'Work on main branch'],
        ['Uncommitted changes from someone else', 'Report the changes, ask developer how to proceed', 'Discard or stash without asking'],
        ['Remote is ahead', 'Pull: git pull origin dev', 'Work on stale code'],
        ['Env vars missing', 'Report which vars are missing, check vault', 'Proceed without required env vars'],
        ['Dev server is down', 'Start it if requested, or report', 'Assume tests will work without server'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. TASK TYPE PLAYBOOKS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Task Type Playbooks', level=1)

para('Each playbook consists of ordered phases. Every phase has a checklist that MUST be completed before moving to the next phase. The agent cannot skip phases unless the developer explicitly says to.')

# ── 7.1 CODING ──────────────────────────────────────────────
doc.add_heading('7.1 CODING', level=2)
para('Goal: Implement or modify code with correctness, minimal blast radius, and full edge case coverage.')

doc.add_heading('7.1.1 Phase 1: Scope & Impact Analysis', level=3)
para('Before writing ANY code, systematically identify everything that will be affected.')

add_checklist([
    'List every file that will be CREATED (new files)',
    'List every file that will be MODIFIED (existing files)',
    'List every database table affected — will schema change? New columns? New tables? New constraints?',
    'List every API route affected — new routes? Modified request/response shapes?',
    'List every UI component that renders the affected data',
    'Identify integration points — does this touch: Twilio, Resend, Google Calendar, Anthropic, Voyage, Kling?',
    'Check: does this overlap with any existing feature? (grep/search codebase BEFORE coding)',
    'Check: does this require a new database migration? If yes, determine the next migration number',
    'Check: does this require new environment variables? If yes, list them',
    'Check: does this require changes to CLAUDE.md, deploy.md, or README?',
    'Check: could this break any existing tests? (read the test files)',
    'Check: is another developer currently working on related files? (check git log --oneline -5 for recent changes)',
])

para('Scenarios to consider during scope analysis:')
add_table(
    ['Scenario', 'Required Action'],
    [
        ['New DB column on existing table', 'Migration file + types/database.ts update + any form that displays the field'],
        ['New DB table', 'Migration + types + server action + API route (if needed) + UI component'],
        ['Modifying a server action', 'Check all callers (grep for the function name) + update types if signature changes'],
        ['Adding a new page', 'Page file + layout check + navigation update + any data fetching'],
        ['Changing an API response shape', 'Update all frontend consumers + types + any external callers'],
        ['Adding an env var', 'Document in CLAUDE.md + add to .env.example + encrypt with vault.sh'],
        ['Modifying a shared component', 'Check every page/component that uses it + verify no visual regressions'],
        ['Changing a Zod schema', 'Check all forms that use it + API routes that validate with it'],
    ]
)

doc.add_heading('7.1.2 Phase 2: Context Loading', level=3)
add_checklist([
    'Read the relevant existing files (NOT the entire repo — only files identified in Phase 1)',
    'Read the relevant type definitions in src/types/database.ts',
    'Read the relevant migration files if touching DB schema',
    'Read the relevant test files if they exist',
    'Read CLAUDE.md coding conventions section',
    'Summarize current behavior in 3-5 bullet points BEFORE modifying anything',
    'If the feature relates to an existing module, understand the data flow end-to-end',
])

doc.add_heading('7.1.3 Phase 3: Implementation Plan', level=3)
add_checklist([
    'Write a short plan: entry points → data flow → new types/functions → error handling',
    'State the error handling strategy explicitly (what errors can occur, how each is handled)',
    'State any breaking changes or migration needs',
    'If plan is complex (5+ files or schema change), present it to the developer BEFORE coding',
    'If plan involves AI/model calls, specify which tier handles what (Haiku/Sonnet/Opus)',
    'Estimate: how many files will change? Is this a 1-commit or multi-commit task?',
])

doc.add_heading('7.1.4 Phase 4: Implementation', level=3)
para('Follow these rules during code writing:')

bold_para('File Organization:')
add_checklist([
    'Create/modify files one at a time — not giant multi-file diffs',
    'Server Actions for mutations → src/actions/',
    'API routes for GETs and external webhooks → src/app/api/',
    'Zod v4 for all validation',
    'JSONB columns for flexible structured data',
    'Path alias: @/ maps to src/',
])

bold_para('UI/Styling:')
add_checklist([
    'Use lf-* CSS classes from the design system (globals.css)',
    'No inline styles — use class names',
    'Emoji icons on pages, Lucide only inside components',
    'force-dynamic on pages that need real-time data',
    'revalidatePath() after every mutation',
])

bold_para('Data Integrity (from feedback_data_integrity.md):')
add_checklist([
    'Validate inputs at ALL boundaries (API, forms, webhooks, server actions)',
    'Use DB constraints: FK, NOT NULL, CHECK, UNIQUE — not just application validation',
    'Use transactions for multi-table mutations',
    'Verify referenced records exist before linking (FK check)',
    'Include rollback/cleanup on partial failures',
    'Test edge cases: duplicates, race conditions, missing references, null values',
    'Parent status NEVER shows "complete" if any child is incomplete (feedback_subtask_completion.md)',
])

bold_para('Security:')
add_checklist([
    'No SQL injection (use parameterized queries via Supabase client)',
    'No XSS (React auto-escapes, but verify any dangerouslySetInnerHTML)',
    'No exposed secrets (never commit .env.local, API keys, tokens)',
    'Validate webhook signatures for inbound webhooks',
    'Sanitize user input before storing in DB',
])

doc.add_heading('7.1.5 Phase 5: Self-Check', level=3)
para('After writing code, BEFORE presenting to the developer:')
add_checklist([
    'Re-read EVERY modified file line by line',
    'Check for: unused variables, unhandled branches, type mismatches',
    'Check for: missing error handling (what if the DB call fails? the API times out? the input is null?)',
    'Check for: missing edge cases (empty array, zero value, max length, special characters)',
    'Run or write tests covering the new behavior',
    'If tests exist: npx vitest run to verify nothing is broken',
    'Verify import paths are correct (no circular dependencies)',
    'Check that async/await is used correctly (no floating promises)',
])

doc.add_heading('7.1.6 Phase 6: Output & Summary', level=3)
add_checklist([
    'Summarize what was built (files changed, behavior added/modified)',
    'List any breaking changes or required manual actions',
    'List any new environment variables',
    'List any new migrations that need to be applied',
    'If tests were written, report: X tests passing out of Y',
    'Suggest: should this be one commit or multiple?',
])

doc.add_heading('7.1.7 Subtype Variations', level=3)
add_table(
    ['Subtype', 'Additional Requirements'],
    [
        ['feature', 'Full Phase 1 scope analysis required. If complex, require Phase 3 plan approval before coding.'],
        ['bugfix', 'Before Phase 4: reproduce the bug, write a FAILING test that captures it. After Phase 4: verify the test now PASSES. Minimal fix only — do NOT refactor surrounding code.'],
        ['refactor', 'State explicitly: "This changes structure, NOT behavior." Existing tests must pass WITHOUT modification. If tests need changes, explain each change.'],
        ['script', 'Make scripts idempotent (safe to run twice). Add dry-run mode. Include clear console output. For seed data: use realistic BC real estate data.'],
    ]
)

doc.add_page_break()

# ── 7.2 TESTING ──────────────────────────────────────────────
doc.add_heading('7.2 TESTING', level=2)
para('Goal: Achieve comprehensive test coverage with deterministic, isolated, meaningful tests.')

doc.add_heading('7.2.1 Phase 1: Clarify Test Level', level=3)
add_checklist([
    'Determine level: unit (pure functions, no I/O), integration (DB/API calls), e2e (full user flow), eval (AI quality)',
    'Identify the specific behaviors to test — not "test everything" but targeted coverage',
    'Check existing test files in app/src/ for patterns and conventions',
])

doc.add_heading('7.2.2 Phase 2: Test Plan — Complete Scenario Coverage', level=3)
para('For EVERY function/feature being tested, cover ALL of the following scenario categories:')

add_table(
    ['Category', 'Scenarios to Cover', 'Example'],
    [
        ['Happy Path', 'Normal operation with valid inputs', 'createContact({name: "John", type: "buyer"}) → returns contact with ID'],
        ['Empty/Null Inputs', 'Missing required fields, null values, undefined, empty strings', 'createContact({}) → throws validation error'],
        ['Boundary Values', 'Min, max, exactly at limit, one above/below limit', 'price = 0, price = 999999999, price = -1'],
        ['Invalid Types', 'Wrong data type, string where number expected, array where object expected', 'createContact({name: 123}) → type error'],
        ['Special Characters', 'Unicode, emoji, HTML entities, SQL injection attempts, XSS payloads', "name = \"O'Brien\", name = \"<script>alert(1)</script>\""],
        ['Duplicate Operations', 'Creating the same record twice, double-submit, concurrent updates', 'Two contacts with same email — what happens?'],
        ['Concurrent Access', 'Two users modifying same record, race conditions', 'Two agents update same contact simultaneously'],
        ['Error Conditions', 'Network failure, DB error, API timeout, service unavailable', 'Supabase returns 500 — does the action handle it?'],
        ['Permission/Auth', 'Unauthenticated request, wrong role, expired token', 'API call without session → 401'],
        ['Cascade Effects', 'Deleting a parent record — what happens to children?', 'Delete contact → what happens to their appointments?'],
        ['Data Integrity', 'Orphaned records, circular references, FK violations', 'Creating an appointment for a non-existent listing'],
        ['State Transitions', 'Invalid transitions, skipping steps, going backwards', 'Moving listing from Phase 1 directly to Phase 5'],
        ['Large Data', 'Many records, long strings, large JSONB objects', '1000 contacts in a list — does pagination work?'],
        ['Time-Dependent', 'Timezone handling, date boundaries, expired data', 'Showing scheduled for yesterday — how is it displayed?'],
    ]
)

doc.add_heading('7.2.3 Phase 3: Implement Tests', level=3)
add_checklist([
    'Use vitest (project standard)',
    'Test files go in app/src/ mirroring the source location',
    'Import from relative paths (e.g., ./chunker not ../../src/lib/rag/chunker)',
    'Tests MUST be deterministic — no dependency on external services, current time, or random values',
    'Tests MUST be isolated — no shared state between tests (no global variables modified)',
    'Use descriptive names: it("returns empty array when contact has no activities")',
    'Group with describe() blocks by function or behavior',
    'For each test: Arrange (setup) → Act (execute) → Assert (verify)',
    'Copy source files to app/src/ test directory if needed for vitest resolution',
])

doc.add_heading('7.2.4 Phase 4: Failure Analysis', level=3)
add_checklist([
    'For each failing test, categorize: environment issue, flaky test, actual bug, or incorrect assertion',
    'Environment issue: fix the environment (missing dep, wrong config), not the test',
    'Flaky test: add retries or fix non-determinism (usually timing or ordering)',
    'Actual bug: the code is wrong — fix the code, keep the test',
    'Incorrect assertion: the test expectation is wrong — fix the test with explanation',
    'Re-run until ALL tests pass — do not leave failing tests',
])

doc.add_heading('7.2.5 Phase 5: Coverage Report', level=3)
add_checklist([
    'State: X tests passing out of Y total',
    'List any skipped tests and why they were skipped',
    'Identify gaps: which scenario categories (from Phase 2 table) are NOT covered?',
    'Recommend: what additional tests would provide the most value?',
])

doc.add_page_break()

# ── 7.3 DEBUGGING ──────────────────────────────────────────
doc.add_heading('7.3 DEBUGGING', level=2)
para('Goal: Systematically find root cause and propose minimal fix. Never guess — investigate.')

doc.add_heading('7.3.1 Phase 1: Problem Statement', level=3)
add_checklist([
    'Restate the symptom precisely: what is happening vs what should happen?',
    'Is there an error message or stack trace? Read it completely — every line.',
    'When does it happen? Always, intermittently, under specific conditions?',
    'What is the scope? All users, one page, one API route, one contact?',
    'Was it working before? If yes, what changed? (check recent git commits)',
    'What environment? Dev, staging, production? (behavior may differ)',
])

doc.add_heading('7.3.2 Phase 2: Reproduction', level=3)
add_checklist([
    'Propose minimal steps to reproduce the issue',
    'If stack trace exists: trace the call path through the code (file → function → line)',
    'If no stack trace: identify the entry point (which page/API/action) and add logging to narrow down',
    'Check: is this a data issue? (bad data in DB causing the error)',
    'Check: is this an environment issue? (missing env var, wrong config)',
    'Check: is this a timing/race condition? (works sometimes, fails sometimes)',
])

doc.add_heading('7.3.3 Phase 3: Hypothesis List', level=3)
add_checklist([
    'Generate 2-4 plausible root causes, ordered by likelihood',
    'For each hypothesis: state what evidence would CONFIRM it and what would ELIMINATE it',
    'Check the most likely hypothesis first — do NOT investigate all in parallel',
    'If first hypothesis is eliminated, move to the next',
    'If all hypotheses fail: broaden the search (check dependencies, infrastructure, external services)',
])

para('Common hypothesis categories:')
add_table(
    ['Category', 'Examples', 'How to Check'],
    [
        ['Data Issue', 'Null where not expected, wrong type, missing FK', 'Query the DB directly, check the specific record'],
        ['Code Logic', 'Wrong condition, off-by-one, missing await', 'Read the code path line by line'],
        ['Type Mismatch', 'String where number expected, undefined property', 'Check TypeScript types vs runtime values'],
        ['Environment', 'Missing env var, wrong URL, expired token', 'Check .env.local, test the connection'],
        ['Race Condition', 'Concurrent writes, stale cache, out-of-order events', 'Add timestamps/logging to track execution order'],
        ['External Service', 'API down, rate limited, changed response format', 'Test the external call in isolation'],
        ['Migration Gap', 'Column missing, constraint wrong, index missing', 'Compare migration files vs actual DB schema'],
    ]
)

doc.add_heading('7.3.4 Phase 4: Fix Proposal', level=3)
add_checklist([
    'Propose the MINIMAL fix — smallest change that resolves the root cause',
    'Explain: "The root cause is X. This fix addresses it by Y."',
    'State trade-offs if any (performance, backward compatibility, complexity)',
    'Do NOT refactor surrounding code — that is a separate CODING:refactor task',
    'Do NOT add "while I\'m here" improvements — scope creep increases risk',
])

doc.add_heading('7.3.5 Phase 5: Regression Prevention', level=3)
add_checklist([
    'Write a test that WOULD HAVE caught this bug (must fail without fix, pass with fix)',
    'Check: are there similar patterns elsewhere? (grep for the same anti-pattern)',
    'If yes: fix all instances (but note them as separate fixes in the commit)',
    'Update any documentation that was misleading about the correct behavior',
])

doc.add_page_break()

# ── 7.4 DESIGN_SPEC ──────────────────────────────────────────
doc.add_heading('7.4 DESIGN_SPEC', level=2)
para('Goal: Produce a thorough design document that enables confident implementation.')

doc.add_heading('7.4.1 Phase 1: Problem Framing', level=3)
add_checklist([
    'Goals: what problem does this solve? Who benefits?',
    'Non-goals: what is explicitly OUT of scope?',
    'Constraints: tech stack (Next.js, Supabase, etc.), performance targets, cost limits, timeline',
    'Success metrics: how do we measure that this works? (quantitative where possible)',
    'Dependencies: what must exist before this can be built?',
])

doc.add_heading('7.4.2 Phase 2: Current State Audit', level=3)
add_checklist([
    'Read relevant existing code, schema, and APIs',
    'Document what already exists that we can REUSE (don\'t rebuild what exists)',
    'Document pain points and limitations of current approach',
    'Document any failed past attempts or prior art in the codebase',
    'Check: what are other developers currently building that might overlap?',
])

doc.add_heading('7.4.3 Phase 3: Design Options', level=3)
add_checklist([
    'Present at LEAST 2 approaches (never just one)',
    'For each approach: pros, cons, risks, estimated complexity (files/migrations/time)',
    'Include cost estimates where applicable (API costs, hosting, storage)',
    'State which approach you recommend and WHY',
    'Consider: build vs buy vs modify-existing',
])

doc.add_heading('7.4.4 Phase 4: Detailed Design', level=3)
add_checklist([
    'Data model: new tables, columns, indexes, constraints, relationships',
    'API surface: new endpoints, request/response schemas, error codes',
    'Component architecture: new files, modified files, component hierarchy',
    'Data flow: how data moves from input → processing → storage → display',
    'State management: client state, server state, cache invalidation',
    'Error handling: what can fail, how each failure is handled',
    'Security: auth requirements, input validation, data access rules',
])

doc.add_heading('7.4.5 Phase 5: Operational Aspects', level=3)
add_checklist([
    'Deployment plan: migrations, env vars, feature flags, rollout order',
    'Monitoring: what to log, what to alert on, what dashboards to check',
    'Failure modes: what happens when X goes down? How to recover?',
    'Performance: expected load, caching strategy, query optimization',
    'Cost: per-query costs for AI/API calls, storage growth, bandwidth',
])

doc.add_heading('7.4.6 Phase 6: Implementation Plan', level=3)
add_checklist([
    'Break into phases: what ships first? what ships later?',
    'File-level detail per phase (which files created/modified)',
    'Test plan per phase',
    'Dependencies between phases (what must complete before what)',
    'Estimated effort per phase',
])

doc.add_page_break()

# ── 7.5 RAG_KB ──────────────────────────────────────────
doc.add_heading('7.5 RAG_KB', level=2)
para('Goal: Design, tune, or extend the RAG system (src/lib/rag/) with measurable quality.')

doc.add_heading('7.5.1 Phase 1: Use Case Clarity', level=3)
add_checklist([
    'What types of questions will users ask?',
    'What data sources feed the answers?',
    'What privacy constraints exist? (no cross-tenant data, no PII in logs)',
    'What freshness requirements? (real-time, daily, weekly)',
    'What accuracy bar? (must cite sources? can generalize?)',
])

doc.add_heading('7.5.2 Phase 2: Content Preparation', level=3)
add_checklist([
    'Define chunking strategy per source (reference src/lib/rag/chunker.ts patterns)',
    'Define metadata schema: what filters are needed? (contact_id, content_type, date_range)',
    'Estimate embedding volume and cost',
    'Check: are the source tables populated with enough data for meaningful retrieval?',
    'Check: is the seed data representative of real usage?',
])

doc.add_heading('7.5.3 Phase 3: Retrieval Configuration', level=3)
add_checklist([
    'Choose search mode: semantic only vs hybrid (semantic + structured filters)',
    'Set top_k per intent (reference src/lib/rag/constants.ts TOP_K_BY_INTENT)',
    'Set similarity threshold (default 0.3, tighten for precision, loosen for recall)',
    'Set context token budget (default 4000)',
    'Define content_type filters per intent',
])

doc.add_heading('7.5.4 Phase 4: Prompting & Grounding', level=3)
add_checklist([
    'System prompt structure: role, tone, rules, citation format',
    'Context window layout: numbered chunks with source metadata',
    'Guardrail patterns: what topics to refuse (legal, tax, financial)',
    'Fallback message when no relevant results found',
    'Voice rules integration from voice_rules table',
])

doc.add_heading('7.5.5 Phase 5: Evaluation', level=3)
add_checklist([
    'Define 20+ test queries covering all intents',
    'For each query: expected content_type matches, expected top result topic',
    'Test: guardrail triggering (legal/tax questions should be blocked)',
    'Test: fallback behavior (query with no matching content)',
    'Test: cross-contact isolation (query about contact A should not return contact B data)',
    'Measure: average similarity score, number of results above threshold',
    'Measure: end-to-end latency (target P95 < 5 seconds)',
])

doc.add_heading('7.5.6 Phase 6: Iteration Plan', level=3)
add_table(
    ['Problem', 'Diagnosis', 'Fix'],
    [
        ['Low precision (irrelevant results)', 'Threshold too low or content_type filter missing', 'Tighten threshold, add type filters, improve chunking'],
        ['Low recall (missing results)', 'Threshold too high, missing data source, poor chunking', 'Lower threshold, add new sources, reduce chunk size'],
        ['Hallucinated facts', 'Grounding rules too weak, context window too small', 'Strengthen system prompt rules, increase top_k'],
        ['Generic answers', 'Chunks too large, losing detail', 'Smaller chunks, more metadata, better composite docs'],
        ['Slow responses', 'Too many chunks, large context, complex queries', 'Reduce top_k, set context token budget, cache embeddings'],
    ]
)

doc.add_page_break()

# ── 7.6 ORCHESTRATION ──────────────────────────────────────
doc.add_heading('7.6 ORCHESTRATION', level=2)
para('Goal: Design or modify AI agent workflows, triggers, and pipeline stages.')

doc.add_heading('7.6.1 Phase 1: Workflow Type', level=3)
add_checklist([
    'Determine pattern: sequential, event-driven, state machine, supervisor, fan-out',
    'Map to existing infrastructure:',
    '  — trigger-engine.ts: CRM event → workflow enrollment',
    '  — workflow-engine.ts: step-by-step execution with delays and conditions',
    '  — contact-evaluator.ts: AI-driven event routing and fan-out',
    '  — trust-gate.ts: trust-level-based routing (ghost/copilot/autonomous)',
    '  — send-governor.ts: rate limiting and sunset logic',
    'Can this be implemented by modifying existing modules? Or does it need new code?',
])

doc.add_heading('7.6.2 Phase 2: States & Transitions', level=3)
add_checklist([
    'Enumerate ALL possible states',
    'For each state: what transitions are allowed?',
    'For each transition: what triggers it? (event, cron, manual action, AI decision)',
    'What are the guard conditions? (must have X before transitioning to Y)',
    'What happens in dead states? (contact stuck, workflow stalled)',
    'How does rollback work? (can we go back to a previous state?)',
])

doc.add_heading('7.6.3 Phase 3: Error Handling', level=3)
add_checklist([
    'Define timeout per step (what if AI call takes 30 seconds?)',
    'Define retry policy: how many retries, backoff strategy, max wait',
    'Define human-in-the-loop breaks: when does the system escalate to the realtor?',
    'Define graceful degradation: what happens if Anthropic API is down? Voyage is down?',
    'Define circuit breaker: after N failures, stop trying and alert',
])

doc.add_heading('7.6.4 Phase 4: Observability', level=3)
add_checklist([
    'Log every decision to agent_decisions table with reasoning',
    'Track latency per stage (embedding, retrieval, generation)',
    'Define alerts for: failed steps, high latency, low quality scores, stuck workflows',
    'Provide admin visibility: which workflows are active, where are contacts stuck?',
])

doc.add_page_break()

# ── 7.7 INTEGRATION ──────────────────────────────────────
doc.add_heading('7.7 INTEGRATION', level=2)
para('Goal: Wire external APIs and services safely and reliably.')

doc.add_heading('7.7.1 Phase 1: API Understanding', level=3)
add_checklist([
    'Read API documentation completely (endpoints, auth mechanism, rate limits, quotas)',
    'Identify: which endpoints do we need? What data do we send/receive?',
    'Identify: is there a sandbox/test environment? (use it for development)',
    'Identify: do we need to receive webhooks? What events? What signature verification?',
    'Check: do we already have a similar integration? (reuse patterns from twilio.ts, resend.ts, etc.)',
])

doc.add_heading('7.7.2 Phase 2: Data Contracts', level=3)
add_checklist([
    'Define request schemas (what we send)',
    'Define response schemas (what we expect back)',
    'Map external fields to internal models (naming differences, type conversions)',
    'Handle: missing fields (API returns less than expected)',
    'Handle: extra fields (API returns more than expected — ignore gracefully)',
    'Handle: field type mismatches (string where number expected)',
])

doc.add_heading('7.7.3 Phase 3: Error & Retry Logic', level=3)
add_checklist([
    'Define timeout values (connection timeout, read timeout)',
    'Implement retry with exponential backoff: 1s → 2s → 4s → give up',
    'Handle rate limiting (429 responses): respect Retry-After header',
    'Ensure idempotency: safe to retry the same request',
    'Handle: 401 (re-authenticate), 403 (permission error), 404 (resource gone), 500 (server error)',
])

doc.add_heading('7.7.4 Phase 4: Security', level=3)
add_checklist([
    'Store API keys in .env.local → encrypt with vault.sh → NEVER commit',
    'Never log API keys, tokens, or PII (mask in error messages)',
    'Validate incoming webhook signatures before processing',
    'Use HTTPS for all external calls',
    'Scope API credentials to minimum required permissions',
])

doc.add_heading('7.7.5 Phase 5: Integration Tests', level=3)
add_table(
    ['Scenario', 'What to Test'],
    [
        ['Happy path', 'Valid request → expected response → data saved correctly'],
        ['Auth failure', '401/403 → graceful error message, no crash'],
        ['Rate limit', '429 → retry with backoff → eventually succeeds or reports'],
        ['Timeout', 'No response in N seconds → timeout error → no hanging process'],
        ['Malformed response', 'API returns unexpected format → handled gracefully'],
        ['Webhook signature', 'Valid signature → processed; invalid signature → rejected 401'],
        ['Network failure', 'DNS failure or connection refused → error message, retry if appropriate'],
    ]
)

doc.add_page_break()

# ── 7.8 DOCS ──────────────────────────────────────────
doc.add_heading('7.8 DOCS', level=2)
para('Goal: Create clear, accurate, maintainable documentation.')

add_checklist([
    'Phase 1 — Audience: Who reads this? Developer, user, or admin? What action does it enable?',
    'Phase 2 — Outline: Structure sections BEFORE writing content',
    'Phase 3 — Draft: Ground examples in actual system (real file paths, real table names, real commands)',
    'Phase 4 — Verify: All file paths exist, all table/column names are current, all commands work',
    'Phase 5 — Consistency: Align terminology with CLAUDE.md, use same naming conventions throughout',
])

# ── 7.9 EVAL ──────────────────────────────────────────
doc.add_heading('7.9 EVAL', level=2)
para('Goal: Define and run evaluations with structured test sets and measurable outcomes.')

add_checklist([
    'Phase 1 — Define success metrics: accuracy, latency, cost, groundedness, user satisfaction',
    'Phase 2 — Create test set: 20+ representative scenarios covering happy path + edge cases',
    'Phase 3 — Scoring method: automatic metrics (regex match, keyword presence) or manual review',
    'Phase 4 — Run & analyze: execute all test cases, record results, identify failure patterns',
    'Phase 5 — Decision: pass/fail per category, recommend ship/iterate/redesign, document trade-offs',
])

# ── 7.10 INFO_QA ──────────────────────────────────────────
doc.add_heading('7.10 INFO_QA', level=2)
para('Goal: Answer questions accurately by referencing the codebase and documentation — no code changes.')

add_checklist([
    'Phase 1 — Restate the question, identify sub-questions',
    'Phase 2 — Research: check codebase (grep/read), CLAUDE.md, memory, git log',
    'Phase 3 — Answer each sub-question explicitly, cite sources (file:line), call out assumptions',
    'Phase 4 — Provide 1-2 concrete examples or edge cases if relevant',
    'Phase 5 — State limitations: what you don\'t know, what information would improve the answer',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. POST-TASK VALIDATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Post-Task Validation Protocol', level=1)

para('After EVERY task completion — no exceptions:')

add_table(
    ['Condition', 'Required Action'],
    [
        ['Code was changed', 'Run: npx vitest run (all tests must pass)'],
        ['Significant changes (5+ files)', 'Run: bash scripts/test-suite.sh (full 73-test suite)'],
        ['Database schema changed', 'Verify migration file: correct number, proper constraints, RLS policy'],
        ['New environment variable added', 'Document in CLAUDE.md, add to .env.example, NEVER commit the value'],
        ['UI component changed', 'Verify it renders correctly (no broken layouts, missing styles)'],
        ['API route changed', 'Verify request/response shapes match types, error codes are correct'],
        ['Any task completed', 'Update todo list: mark completed, add discovered follow-ups'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. MODEL CHAINING
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Model Chaining Strategy', level=1)

add_table(
    ['Task Phase', 'Model', 'Why', 'Cost'],
    [
        ['Task classification', 'Haiku (internal reasoning)', 'Fast routing, no API cost', 'Free'],
        ['File search, codebase exploration', 'Haiku sub-agents', 'Quick lookup, parallel search', '~$0.001'],
        ['Code implementation', 'Sonnet sub-agents', 'Good balance of speed + code quality', '~$0.01'],
        ['Architecture decisions, complex analysis', 'Opus (main agent)', 'Deep reasoning, multi-file understanding', '~$0.05'],
        ['Test writing', 'Sonnet sub-agents', 'Straightforward but needs accuracy', '~$0.01'],
        ['Code review / self-check', 'Opus (main agent)', 'Catches subtle logic bugs', '~$0.05'],
        ['Design spec writing', 'Opus (main agent)', 'Requires holistic system understanding', '~$0.05'],
        ['Documentation writing', 'Sonnet sub-agents', 'Content generation, not deep reasoning', '~$0.01'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. ASK vs ACT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. Decision Framework: Ask vs Act', level=1)

section('10.1 Act Immediately (Full Permission Granted)')
add_checklist([
    'Request is clear and scoped',
    'Change is localized (single file, single function)',
    'Tests exist and will validate the change',
    'Standard patterns being followed (from CLAUDE.md)',
    'File reads, codebase searches, git operations',
    'Running tests, builds, health checks',
    'Creating migrations, seed data, test files',
])

section('10.2 Ask First')
add_checklist([
    'Request is ambiguous (could mean 2+ different things)',
    'Change affects shared state (DB schema, API contracts, environment variables)',
    'Change could break other developers\' in-progress work',
    'Task classification confidence is LOW',
    'Destructive operation requested (delete files, drop tables, force push)',
    'Change involves 10+ files — present plan first',
    'Trade-off exists with no clear winner — present options',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. EXISTING INFRASTRUCTURE MAPPING
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. Existing Infrastructure Mapping', level=1)

para('The playbook REFERENCES these existing components. It does NOT modify or replace them.')

add_table(
    ['Component', 'Location', 'Playbook Relationship'],
    [
        ['CLAUDE.md', 'Repo root', 'Referenced for coding conventions, tech stack, project structure'],
        ['Memory files', '~/.claude/projects/.../memory/', 'Referenced for behavioral rules (health check, data integrity, etc.)'],
        ['Deploy skill', '.claude/skills/deploy.md', 'Called by playbook for deployment tasks'],
        ['Test skill', '.claude/skills/test.md', 'Called by playbook for validation'],
        ['Health check script', 'scripts/health-check.sh', 'Called by pre-flight protocol'],
        ['Test suite script', 'scripts/test-suite.sh', 'Called by post-task validation'],
        ['AI agent modules', 'src/lib/ai-agent/', 'NOT modified by playbook — orchestration rules are separate'],
        ['RAG system', 'src/lib/rag/', 'NOT modified by playbook — RAG_KB tasks extend it'],
        ['Workflow engine', 'src/lib/workflow-engine.ts', 'NOT modified by playbook — ORCHESTRATION tasks extend it'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 12. NON-DISRUPTION GUARANTEES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('12. Integration Points & Non-Disruption Guarantees', level=1)

section('12.1 What the Playbook Changes')
add_checklist([
    'ONE new file: .claude/agent-playbook.md (the playbook itself)',
    'No application code changes',
    'No database schema changes',
    'No API route changes',
    'No UI changes',
    'No dependency changes (no npm install)',
    'No environment variable changes',
])

section('12.2 What the Playbook Does NOT Change')
add_checklist([
    'CLAUDE.md — referenced, not modified',
    'Memory files — referenced, not modified',
    'Existing skills — referenced, not modified',
    'Health check and test scripts — called, not modified',
    'Any source code in src/ — untouched',
    'Any migration files — untouched',
    'Package.json — untouched',
    'Git history — no rewrites, no force pushes',
])

section('12.3 Multi-Developer Safety')
add_checklist([
    'Pre-flight checks for uncommitted changes from other developers',
    'Agent pulls latest before starting work',
    'Agent does not discard or stash other developers\' changes',
    'Agent uses atomic commits (one concern per commit)',
    'Agent does not amend existing commits (creates new ones)',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 13. TEST SCENARIOS FOR THE PLAYBOOK
# ═══════════════════════════════════════════════════════════════
doc.add_heading('13. Test Scenarios for the Playbook Itself', level=1)

para('These scenarios verify that the agent correctly follows the playbook:')

add_table(
    ['#', 'Scenario', 'Input', 'Expected Agent Behavior'],
    [
        ['T1', 'Simple feature request', '"Add a phone_type field to contacts"', 'Classifies as CODING:feature. Runs full 6-phase checklist. Creates migration + type + form + action.'],
        ['T2', 'Bug report with stack trace', '"Getting TypeError: cannot read property name of undefined on contact page"', 'Classifies as DEBUGGING:error. Does NOT start coding. Traces stack, generates hypotheses, proposes minimal fix.'],
        ['T3', 'Ambiguous request', '"Make the newsletter better"', 'Confidence LOW. Asks: "Better in what way — content quality, design, send timing, or deliverability?"'],
        ['T4', 'Design before build', '"We need to add offer management to the CRM"', 'Classifies as DESIGN_SPEC:feature_spec. Produces design doc first. Does NOT start coding.'],
        ['T5', 'Test writing request', '"Write tests for the RAG chunker"', 'Classifies as TESTING:unit. Covers ALL scenario categories from Phase 2 table. Produces ≥15 tests.'],
        ['T6', 'Refactor request', '"Clean up the contact form component"', 'Classifies as CODING:refactor. States "changes structure NOT behavior". Existing tests must pass unchanged.'],
        ['T7', 'Multi-type request', '"Add a webhook for Stripe and write tests"', 'Primary: INTEGRATION:webhook. Secondary: TESTING:integration. Completes integration checklist first, then testing.'],
        ['T8', 'Info question', '"How does the trust gate system work?"', 'Classifies as INFO_QA:explain. Reads trust-gate.ts and trust-manager.ts. Answers with file:line citations.'],
        ['T9', 'Pre-flight failure', 'Health check shows build failure', 'Reports failure. Offers to fix build. Does NOT proceed with original task.'],
        ['T10', 'Another dev\'s uncommitted changes', 'git status shows modified files not by this session', 'Reports changes. Asks developer how to proceed. Does NOT discard.'],
        ['T11', 'Guardrail: main branch', 'Agent finds itself on main branch', 'Switches to dev. Does NOT work on main.'],
        ['T12', 'Complex feature (5+ files)', '"Build a mortgage calculator page"', 'Classifies as CODING:feature. After Phase 3 plan: presents plan to developer BEFORE coding.'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 14. APPENDIX: QUICK REFERENCE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('14. Appendix: Quick Reference Card', level=1)

para('Compact reference for the agent to check before every task:')

bold_para('PRE-FLIGHT')
para('□ health-check.sh  □ branch=dev  □ pull latest  □ load memory')

bold_para('CLASSIFY')
para('□ task_type:subtype  □ confidence  □ reasoning  □ affected files')

bold_para('EXECUTE (per task type)')
para('''CODING: scope → context → plan → implement → self-check → output
TESTING: level → scenarios (ALL categories) → implement → analyze → report
DEBUGGING: problem → reproduce → hypotheses → inspect → fix → regression test
DESIGN: problem → audit → options → design → ops → plan
RAG_KB: use case → content → retrieval → prompts → eval → iterate
ORCHESTRATION: type → states → errors → observability
INTEGRATION: API docs → contracts → retry → security → tests
DOCS: audience → outline → draft → verify → consistency
EVAL: metrics → test set → scoring → run → decision
INFO_QA: restate → research → answer → examples → limitations''')

bold_para('VALIDATE')
para('□ vitest run  □ test-suite.sh (if major)  □ no regressions  □ update todos')

bold_para('MODEL CHAIN')
para('Haiku=search/classify  Sonnet=code/tests/docs  Opus=architecture/review/debug')

# ── Save ──────────────────────────────────────────────────
script_dir = os.path.dirname(os.path.abspath(__file__))
repo_root = os.path.dirname(script_dir)
output_path = os.path.join(repo_root, 'docs', 'SPEC_Agent_Playbook.docx')
doc.save(output_path)
print(f"✅ Saved to: {output_path}")
print(f"   Pages: ~{len(doc.paragraphs) // 40} (estimated)")
print(f"   Sections: 14")
print(f"   Tables: {len(doc.tables)}")
print(f"   Checklists: ~{sum(1 for p in doc.paragraphs if p.style.name == 'List Bullet')} items")
